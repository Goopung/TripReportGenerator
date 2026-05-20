import io
import json
import os
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas
from pypdf import PdfReader, PdfWriter

from core.data_model import ReasonStatementData, TripReportData
from core.file_utils import IMAGE_EXTS, PDF_EXTS, collect_paths


def _maybe_pdf_to_images(path: str | Path, temp_dir: str | Path, max_pages: int = 3) -> list[str]:
    path = Path(path)
    temp_dir = Path(temp_dir)
    temp_dir.mkdir(parents=True, exist_ok=True)

    if path.suffix.lower() not in PDF_EXTS:
        return []

    try:
        import fitz

        doc = fitz.open(str(path))
        out = []
        for idx in range(min(max_pages, len(doc))):
            page = doc[idx]
            pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
            img_path = temp_dir / f"{path.stem}_page_{idx + 1}.png"
            pix.save(str(img_path))
            out.append(str(img_path))
        return out
    except Exception:
        return []


def _paths_to_insertable_images(paths: list[str], temp_dir: str | Path) -> list[str]:
    images: list[str] = []
    for path in paths:
        p = Path(path)
        if p.suffix.lower() in IMAGE_EXTS:
            images.append(str(p))
        elif p.suffix.lower() in PDF_EXTS:
            images.extend(_maybe_pdf_to_images(p, temp_dir))
    return images


def _set_doc_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Malgun Gothic"


def _para(doc: Document, text: str = "", bold: bool = False, align: Any = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Malgun Gothic"
    run.font.size = Pt(10.5)


def _add_image_list_docx(doc: Document, paths: list[str], temp_dir: str | Path, max_width_inches: float = 6.1) -> None:
    image_paths = _paths_to_insertable_images(paths, temp_dir)
    for image_path in image_paths:
        try:
            doc.add_picture(image_path, width=Inches(max_width_inches))
            doc.add_paragraph("")
        except Exception:
            continue


def _dot_date(value: str) -> str:
    try:
        from datetime import datetime

        d = datetime.fromisoformat(str(value)).date()
        return f"{d.year}. {d.month}. {d.day}"
    except Exception:
        return str(value)


def _month_day(value: str) -> str:
    try:
        from datetime import datetime

        d = datetime.fromisoformat(str(value)).date()
        return f"{d.month}월 {d.day}일"
    except Exception:
        return str(value)


def _date_range_label(start: str, end: str) -> str:
    return f"{_dot_date(start)}. ~ {_dot_date(end)}."


def _section_docx(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"□ {title}")
    r.bold = True
    r.font.name = "Batang"
    r.font.size = Pt(13)


def _bullet_docx(doc: Document, text: str, level: int = 0) -> None:
    prefix = "o " if level == 0 else "- "
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(prefix + str(text).strip())
    r.font.name = "Batang"
    r.font.size = Pt(10.5)


def _clean_expected_effect_line(text: str) -> str:
    text = str(text).strip()
    text = re.sub(r"^\s*[-•●*]\s*", "", text)
    text = re.sub(r"^\s*\d+[\).]\s*", "", text)
    return text.strip()


def _expected_effect_paragraphs(data: TripReportData) -> list[str]:
    expected_effect_text = getattr(data, "expected_effect_text", "").strip()

    if not expected_effect_text:
        expected_effect_text = (
            f"이번 {data.conference_name} 참여를 통해 파악한 최신 연구 동향과 주요 세션 내용은 "
            f"향후 본 연구과제의 연구 방향을 구체화하고 관련 분야의 학술적 네트워크를 확대하는 데 "
            f"기여할 것으로 기대된다."
        )

    paragraphs = []
    for line in expected_effect_text.split("\n"):
        cleaned = _clean_expected_effect_line(line)
        if cleaned:
            paragraphs.append(cleaned)

    if not paragraphs:
        paragraphs.append(
            f"이번 {data.conference_name} 참여를 통해 파악한 최신 연구 동향과 주요 세션 내용은 "
            f"향후 본 연구과제의 연구 방향을 구체화하고 관련 분야의 학술적 네트워크를 확대하는 데 "
            f"기여할 것으로 기대된다."
        )

    return paragraphs


def _set_docx_cell(cell, text: str, bold: bool = False, shade: str | None = None, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    if shade:
        _shade_cell(cell, shade)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Batang"
    r.font.size = Pt(9.5)


def _add_docx_image_grid(doc: Document, paths: list[str], temp_dir: str | Path, cols: int = 2, width_inches: float = 2.7) -> None:
    image_paths = _paths_to_insertable_images(paths, temp_dir)
    if not image_paths:
        return

    table = doc.add_table(rows=0, cols=cols)
    table.alignment = 1

    for idx in range(0, len(image_paths), cols):
        row_cells = table.add_row().cells
        for col in range(cols):
            img_idx = idx + col
            if img_idx >= len(image_paths):
                continue
            p = row_cells[col].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                run = p.add_run()
                run.add_picture(image_paths[img_idx], width=Inches(width_inches))
            except Exception:
                pass

    doc.add_paragraph("")


def generate_trip_docx(data: TripReportData, output_path: str | Path) -> str:
    output_path = Path(output_path)
    temp_dir = output_path.parent / "_temp_images_docx"
    temp_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_doc_style(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    doc.add_paragraph("\n\n\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{data.conference_name} 학회 참석")
    r.bold = True
    r.font.name = "Batang"
    r.font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{data.trip_type}출장 결과보고서")
    r.bold = True
    r.font.name = "Batang"
    r.font.size = Pt(24)

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(_dot_date(data.report_date))
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)

    doc.add_page_break()

    _section_docx(doc, "출장목적")
    for paragraph in [x.strip() for x in str(data.purpose_text).split("\n") if x.strip()]:
        _bullet_docx(doc, paragraph, 0)

    _section_docx(doc, "출장기간")
    _bullet_docx(doc, _date_range_label(data.start_date, data.end_date), 0)

    doc.add_page_break()

    _section_docx(doc, "출장인")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["출장자", "출장기간", "출장지", "학회장소"]
    for i, h in enumerate(headers):
        _set_docx_cell(table.rows[0].cells[i], h, bold=True, shade="D9D9D9")

    row = table.add_row().cells
    _set_docx_cell(row[0], data.traveler_name)
    _set_docx_cell(row[1], _date_range_label(data.start_date, data.end_date))
    _set_docx_cell(row[2], data.destination)
    _set_docx_cell(row[3], data.venue)

    _section_docx(doc, "세부일정")
    schedule_table = doc.add_table(rows=1, cols=3)
    schedule_table.style = "Table Grid"
    headers = ["일자", "학회장소", "세부일정"]
    for i, h in enumerate(headers):
        _set_docx_cell(schedule_table.rows[0].cells[i], h, bold=True, shade="D9D9D9")

    for row_data in data.daily_schedule:
        cells = schedule_table.add_row().cells
        _set_docx_cell(cells[0], _month_day(row_data.get("date", "")))
        _set_docx_cell(cells[1], data.venue)
        _set_docx_cell(cells[2], row_data.get("content", ""), align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_page_break()

    _section_docx(doc, "본 연구와 관련성 및 주요 세션 요약")
    for paragraph in [x.strip() for x in str(data.research_relatedness).split("\n") if x.strip()]:
        if paragraph.startswith(("1.", "2.", "3.", "4.", "5.", "-", "*", "ü")):
            _bullet_docx(doc, paragraph, 1)
        else:
            _bullet_docx(doc, paragraph, 0)

    doc.add_page_break()

    _section_docx(doc, "증빙(탑승권/기타증빙) 및 현장 방문 사진")
    if data.traveler_name:
        _bullet_docx(doc, data.traveler_name, 0)

    transport_paths = []
    for key in ["boarding_pass", "e_ticket", "ticket_receipt", "acceptance_letter"]:
        transport_paths.extend(data.transport_files.get(key, []))
    if transport_paths:
        _bullet_docx(doc, "탑승권 및 항공권 관련 증빙", 0)
        _add_docx_image_grid(doc, transport_paths, temp_dir, cols=2, width_inches=2.8)

    all_receipts = []
    for row in data.daily_schedule:
        all_receipts.extend(data.daily_receipts.get(str(row.get("date", "")), []))
    if all_receipts:
        _bullet_docx(doc, "영수증", 0)
        _add_docx_image_grid(doc, all_receipts, temp_dir, cols=3, width_inches=1.9)

    doc.add_page_break()

    _bullet_docx(doc, f"{data.conference_name} 학회 참석 사진", 0)
    for row in data.daily_schedule:
        date_key = str(row.get("date", ""))
        paths = data.daily_photos.get(date_key, [])
        if not paths:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(_month_day(date_key))
        r.font.name = "Batang"
        r.font.size = Pt(12)
        _add_docx_image_grid(doc, paths, temp_dir, cols=2, width_inches=2.8)

    lodging_paths = []
    lodging_paths.extend(data.lodging_files.get("lodging_confirmation", []))
    lodging_paths.extend(data.lodging_files.get("lodging_receipt", []))
    if lodging_paths:
        doc.add_page_break()
        _bullet_docx(doc, "숙박 관련 증빙자료", 0)
        _add_docx_image_grid(doc, lodging_paths, temp_dir, cols=2, width_inches=2.8)

    extra_paths = []
    for value in data.extra_files.values():
        extra_paths.extend(value)
    if extra_paths:
        doc.add_page_break()
        _bullet_docx(doc, "기타 증빙자료", 0)
        _add_docx_image_grid(doc, extra_paths, temp_dir, cols=2, width_inches=2.8)

    doc.add_page_break()
    _section_docx(doc, "기대효과")
    for paragraph in _expected_effect_paragraphs(data):
        _bullet_docx(
            doc,
            f"{paragraph}",
            0,
        )

    doc.save(str(output_path))
    return str(output_path)


def _font_candidates(kind: str) -> list[str]:
    if kind == "korean":
        env = os.getenv("KOREAN_FONT_PATH", "").strip()
        candidates = [env] if env else []
        candidates.extend([
            "C:/Windows/Fonts/batang.ttc",
            "C:/Windows/Fonts/batang.ttf",
            "C:/Windows/Fonts/malgun.ttf",
            "C:/Windows/Fonts/gulim.ttc",
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/Library/Fonts/AppleGothic.ttf",
            "/System/Library/Fonts/Supplemental/AppleGothic.ttf",
            "/usr/share/fonts/opentype/noto/NotoSerifCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSerifCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/nanum/NanumMyeongjo.ttf",
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        ])
        return [c for c in candidates if c]

    env = os.getenv("ENGLISH_FONT_PATH", "").strip()
    candidates = [env] if env else []
    candidates.extend([
        "C:/Windows/Fonts/times.ttf",
        "C:/Windows/Fonts/timesbd.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSerif-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
    ])
    return [c for c in candidates if c]


def _register_font(name: str, candidates: list[str], cid_fallback: str | None = None, fallback: str = "Helvetica") -> str:
    for font_path in candidates:
        if font_path and Path(font_path).exists():
            try:
                pdfmetrics.registerFont(TTFont(name, font_path))
                return name
            except Exception:
                continue

    if cid_fallback:
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(cid_fallback))
            return cid_fallback
        except Exception:
            pass

    return fallback


def _register_pdf_fonts() -> tuple[str, str]:
    korean_font = _register_font(
        "KoreanBatang",
        _font_candidates("korean"),
        cid_fallback="HYSMyeongJo-Medium",
        fallback="Helvetica",
    )
    english_font = _register_font(
        "EnglishTimes",
        _font_candidates("english"),
        fallback="Times-Roman",
    )
    return korean_font, english_font


def _is_korean_char(ch: str) -> bool:
    code = ord(ch)
    return (
        0xAC00 <= code <= 0xD7A3 or
        0x1100 <= code <= 0x11FF or
        0x3130 <= code <= 0x318F or
        0x4E00 <= code <= 0x9FFF
    )


def _pick_font(ch: str, korean_font: str, english_font: str) -> str:
    if _is_korean_char(ch):
        return korean_font
    if ch.isascii():
        return english_font
    return korean_font


def _text_width(text: str, font_size: float, korean_font: str, english_font: str) -> float:
    total = 0.0
    for ch in str(text):
        total += pdfmetrics.stringWidth(ch, _pick_font(ch, korean_font, english_font), font_size)
    return total


def _wrap_mixed_text(text: str, max_width: float, font_size: float, korean_font: str, english_font: str) -> list[str]:
    words = re.split(r"(\s+)", str(text).replace("\r", ""))
    lines: list[str] = []
    current = ""

    def too_wide(s: str) -> bool:
        return _text_width(s, font_size, korean_font, english_font) > max_width

    for part in words:
        if part == "":
            continue
        if "\n" in part:
            pieces = part.split("\n")
            for idx, piece in enumerate(pieces):
                if piece:
                    candidate = current + piece
                    if current and too_wide(candidate):
                        lines.append(current.rstrip())
                        current = piece
                    else:
                        current = candidate
                if idx < len(pieces) - 1:
                    lines.append(current.rstrip())
                    current = ""
            continue

        candidate = current + part
        if current and too_wide(candidate):
            lines.append(current.rstrip())
            current = part.lstrip()
        else:
            current = candidate

        while current and too_wide(current):
            cut = ""
            rest = current
            for j, ch in enumerate(current):
                if _text_width(cut + ch, font_size, korean_font, english_font) <= max_width:
                    cut += ch
                else:
                    rest = current[j:]
                    break
            if cut:
                lines.append(cut.rstrip())
                current = rest.lstrip()
            else:
                break

    if current.strip():
        lines.append(current.rstrip())
    return lines


def _draw_mixed_line(c: canvas.Canvas, text: str, x: float, y: float, font_size: float, korean_font: str, english_font: str) -> None:
    cursor = x
    for ch in str(text):
        font = _pick_font(ch, korean_font, english_font)
        c.setFont(font, font_size)
        c.drawString(cursor, y, ch)
        cursor += pdfmetrics.stringWidth(ch, font, font_size)


def _draw_mixed_paragraph(
    c: canvas.Canvas,
    text: str,
    x: float,
    y: float,
    max_width: float,
    max_height: float,
    font_size: float,
    leading: float,
    korean_font: str,
    english_font: str,
) -> None:
    lines = _wrap_mixed_text(text, max_width, font_size, korean_font, english_font)
    max_lines = max(int(max_height // leading), 1)
    for i, line in enumerate(lines[:max_lines]):
        _draw_mixed_line(c, line, x, y - i * leading, font_size, korean_font, english_font)


def _pdf_style():
    base = getSampleStyleSheet()
    korean_font, english_font = _register_pdf_fonts()
    font = korean_font
    return {
        "font_name": font,
        "korean_font": korean_font,
        "english_font": english_font,
        "title": ParagraphStyle("title", parent=base["Title"], fontName=font, fontSize=24, alignment=TA_CENTER, leading=34),
        "cover_subtitle": ParagraphStyle("cover_subtitle", parent=base["Title"], fontName=font, fontSize=26, alignment=TA_CENTER, leading=36),
        "h1": ParagraphStyle("h1", parent=base["Heading1"], fontName=font, fontSize=15, leading=21, spaceBefore=4, spaceAfter=8),
        "h2": ParagraphStyle("h2", parent=base["Heading2"], fontName=font, fontSize=12, leading=17, spaceBefore=8, spaceAfter=5),
        "normal": ParagraphStyle("normal", parent=base["Normal"], fontName=font, fontSize=10, leading=15),
        "bullet": ParagraphStyle("bullet", parent=base["Normal"], fontName=font, fontSize=10, leading=15, leftIndent=10),
        "center": ParagraphStyle("center", parent=base["Normal"], fontName=font, fontSize=12, alignment=TA_CENTER, leading=18),
        "small_center": ParagraphStyle("small_center", parent=base["Normal"], fontName=font, fontSize=10, alignment=TA_CENTER, leading=14),
    }


def _p(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def _cell(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(str(text).replace("\n", "<br/>"), style)


def _section_pdf(story: list, title: str, styles: dict) -> None:
    story.append(_p(f"□ {title}", styles["h1"]))


def _bullet_pdf(story: list, text: str, styles: dict, level: int = 0) -> None:
    prefix = "● " if level == 0 else "- "
    story.append(_p(prefix + str(text).strip(), styles["bullet"]))


def _add_image_list_pdf(story: list, paths: list[str], temp_dir: str | Path) -> None:
    _add_image_grid_pdf(story, paths, temp_dir, cols=2)


def _add_image_grid_pdf(story: list, paths: list[str], temp_dir: str | Path, cols: int = 2) -> None:
    image_paths = _paths_to_insertable_images(paths, temp_dir)
    if not image_paths:
        return

    page_width = 160 * mm
    gap = 5 * mm
    cell_w = (page_width - gap * (cols - 1)) / cols
    max_h = 85 * mm if cols == 2 else 70 * mm

    rows = []
    for i in range(0, len(image_paths), cols):
        row = []
        for j in range(cols):
            idx = i + j
            if idx >= len(image_paths):
                row.append("")
                continue
            try:
                img = Image(image_paths[idx])
                ratio = min(cell_w / img.drawWidth, max_h / img.drawHeight, 1)
                img.drawWidth *= ratio
                img.drawHeight *= ratio
                row.append(img)
            except Exception:
                row.append("")
        rows.append(row)

    table = Table(rows, colWidths=[cell_w] * cols)
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    story.append(Spacer(1, 5 * mm))


def _draw_page_number(c, doc):
    if doc.page <= 1:
        return
    try:
        korean_font, _ = _register_pdf_fonts()
        c.setFont(korean_font, 9)
    except Exception:
        c.setFont("Helvetica", 9)
    c.drawCentredString(A4[0] / 2, 10 * mm, f"- {doc.page} -")


def generate_trip_pdf(data: TripReportData, output_path: str | Path) -> str:
    output_path = Path(output_path)
    temp_dir = output_path.parent / "_temp_images_pdf"
    temp_dir.mkdir(parents=True, exist_ok=True)

    styles = _pdf_style()
    table_font = styles["font_name"]

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story: list = []

    story.append(Spacer(1, 78 * mm))
    story.append(_p(f"{data.conference_name} 학회 참석", styles["title"]))
    story.append(Spacer(1, 4 * mm))
    story.append(_p(f"{data.trip_type}출장 결과보고서", styles["cover_subtitle"]))
    story.append(Spacer(1, 20 * mm))
    story.append(_p(_dot_date(data.report_date), styles["center"]))
    story.append(PageBreak())

    _section_pdf(story, "출장목적", styles)
    for paragraph in [x.strip() for x in str(data.purpose_text).split("\n") if x.strip()]:
        _bullet_pdf(story, paragraph, styles, 0)

    story.append(Spacer(1, 5 * mm))
    _section_pdf(story, "출장기간", styles)
    _bullet_pdf(story, _date_range_label(data.start_date, data.end_date), styles, 0)
    story.append(PageBreak())

    _section_pdf(story, "출장인", styles)
    overview_rows = [
        [_cell("출장자", styles["small_center"]), _cell("출장기간", styles["small_center"]), _cell("출장지", styles["small_center"]), _cell("학회장소", styles["small_center"])],
        [_cell(data.traveler_name, styles["small_center"]), _cell(_date_range_label(data.start_date, data.end_date), styles["small_center"]), _cell(data.destination, styles["small_center"]), _cell(data.venue, styles["small_center"])],
    ]
    t = Table(overview_rows, colWidths=[32 * mm, 45 * mm, 37 * mm, 48 * mm])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, -1), table_font),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t)
    story.append(Spacer(1, 8 * mm))

    _section_pdf(story, "세부일정", styles)
    rows = [[_cell("일자", styles["small_center"]), _cell("학회장소", styles["small_center"]), _cell("세부일정", styles["small_center"])]]
    for row in data.daily_schedule:
        rows.append([
            _cell(_month_day(row.get("date", "")), styles["small_center"]),
            _cell(data.venue, styles["small_center"]),
            _cell(row.get("content", ""), styles["normal"]),
        ])
    t = Table(rows, colWidths=[30 * mm, 45 * mm, 87 * mm])
    t.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("FONTNAME", (0, 0), (-1, -1), table_font),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (1, -1), "CENTER"),
        ("ALIGN", (2, 1), (2, -1), "LEFT"),
    ]))
    story.append(t)
    story.append(PageBreak())

    _section_pdf(story, "본 연구와 관련성 및 주요 세션 요약", styles)
    for paragraph in [x.strip() for x in str(data.research_relatedness).split("\n") if x.strip()]:
        if paragraph.startswith(("1.", "2.", "3.", "4.", "5.", "-", "*", "ü")):
            _bullet_pdf(story, paragraph, styles, 1)
        else:
            _bullet_pdf(story, paragraph, styles, 0)
    story.append(PageBreak())

    _section_pdf(story, "증빙(탑승권/기타증빙) 및 현장 방문 사진", styles)
    if data.traveler_name:
        story.append(_p(data.traveler_name, styles["normal"]))

    transport_paths = []
    for key in ["boarding_pass", "e_ticket", "ticket_receipt", "acceptance_letter"]:
        transport_paths.extend(data.transport_files.get(key, []))
    if transport_paths:
        _bullet_pdf(story, "탑승권 및 항공권 관련 증빙", styles, 0)
        _add_image_grid_pdf(story, transport_paths, temp_dir, cols=2)

    all_receipts = []
    for row in data.daily_schedule:
        all_receipts.extend(data.daily_receipts.get(str(row.get("date", "")), []))
    if all_receipts:
        _bullet_pdf(story, "영수증", styles, 0)
        _add_image_grid_pdf(story, all_receipts, temp_dir, cols=3)

    story.append(PageBreak())

    _bullet_pdf(story, f"{data.conference_name} 학회 참석 사진", styles, 0)
    for row in data.daily_schedule:
        date_key = str(row.get("date", ""))
        paths = data.daily_photos.get(date_key, [])
        if not paths:
            continue
        story.append(_p(_month_day(date_key), styles["center"]))
        _add_image_grid_pdf(story, paths, temp_dir, cols=2)

    lodging_paths = []
    lodging_paths.extend(data.lodging_files.get("lodging_confirmation", []))
    lodging_paths.extend(data.lodging_files.get("lodging_receipt", []))
    if lodging_paths:
        story.append(PageBreak())
        _bullet_pdf(story, "숙박 관련 증빙자료", styles, 0)
        _add_image_grid_pdf(story, lodging_paths, temp_dir, cols=2)

    extra_paths = []
    for value in data.extra_files.values():
        extra_paths.extend(value)
    if extra_paths:
        story.append(PageBreak())
        _bullet_pdf(story, "기타 증빙자료", styles, 0)
        _add_image_grid_pdf(story, extra_paths, temp_dir, cols=2)

    story.append(PageBreak())
    _section_pdf(story, "기대효과", styles)
    for paragraph in _expected_effect_paragraphs(data):
        _bullet_pdf(
            story,
            f"{paragraph}",
            styles,
            0,
        )

    doc.build(story, onFirstPage=_draw_page_number, onLaterPages=_draw_page_number)
    return str(output_path)


def _shade_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def generate_reason_docx(reason: ReasonStatementData, output_path: str | Path) -> str:
    output_path = Path(output_path)
    doc = Document()
    _set_doc_style(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("사  유  서")
    r.bold = True
    r.underline = True
    r.font.name = "Malgun Gothic"
    r.font.size = Pt(20)

    doc.add_paragraph("")

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    rows = [
        ("연구책임자 소속", reason.principal_affiliation, "연구책임자 성명", reason.principal_name),
        ("지원기관", reason.funding_agency, "당해연도 연구기간", reason.research_period),
        ("연구과제명", reason.project_title, "", ""),
    ]
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
            if j in [0, 2]:
                _shade_cell(table.cell(i, j), "D9D9D9")
        if i == 2:
            table.cell(i, 1).merge(table.cell(i, 3))

    doc.add_paragraph("")

    table2 = doc.add_table(rows=2, cols=2)
    table2.style = "Table Grid"
    table2.cell(0, 0).text = "제 목"
    table2.cell(0, 1).text = reason.generated_title
    table2.cell(1, 0).text = "내 용"
    table2.cell(1, 1).text = reason.generated_content
    _shade_cell(table2.cell(0, 0), "D9D9D9")
    _shade_cell(table2.cell(1, 0), "D9D9D9")

    doc.save(str(output_path))
    return str(output_path)


def generate_reason_pdf_on_template(
    reason: ReasonStatementData,
    output_path: str | Path,
    template_path: str | Path,
    reason_date: str,
) -> str:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    template_path = Path(template_path)

    korean_font, english_font = _register_pdf_fonts()

    reader = PdfReader(str(template_path))
    page = reader.pages[0]
    w = float(page.mediabox.width)
    h = float(page.mediabox.height)

    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=(w, h))

    field_font = 9.5
    _draw_mixed_line(c, reason.principal_affiliation, 132, 741, field_font, korean_font, english_font)
    _draw_mixed_line(c, reason.principal_name, 390, 741, field_font, korean_font, english_font)
    _draw_mixed_line(c, reason.funding_agency, 132, 719, field_font, korean_font, english_font)
    _draw_mixed_line(c, reason.research_period, 390, 719, field_font, korean_font, english_font)
    _draw_mixed_line(c, reason.project_title, 132, 696, field_font, korean_font, english_font)

    title_font = 10.5
    _draw_mixed_paragraph(
        c,
        reason.generated_title,
        x=132,
        y=659,
        max_width=410,
        max_height=20,
        font_size=title_font,
        leading=13,
        korean_font=korean_font,
        english_font=english_font,
    )

    content_font = 10.2
    _draw_mixed_paragraph(
        c,
        reason.generated_content,
        x=132,
        y=630,
        max_width=410,
        max_height=405,
        font_size=content_font,
        leading=15,
        korean_font=korean_font,
        english_font=english_font,
    )

    try:
        from datetime import datetime

        parsed = datetime.fromisoformat(str(reason_date)).date()
        year = str(parsed.year)
        month = f"{parsed.month:02d}"
        day = f"{parsed.day:02d}"
    except Exception:
        parts = re.findall(r"\d+", str(reason_date))
        year = parts[0] if len(parts) > 0 else ""
        month = parts[1] if len(parts) > 1 else ""
        day = parts[2] if len(parts) > 2 else ""

    _draw_mixed_line(c, year, 445, 162, 10, korean_font, english_font)
    _draw_mixed_line(c, month, 485, 162, 10, korean_font, english_font)
    _draw_mixed_line(c, day, 509, 162, 10, korean_font, english_font)

    _draw_mixed_line(c, reason.principal_name, 452, 140, 10, korean_font, english_font)

    c.save()
    packet.seek(0)

    overlay_reader = PdfReader(packet)
    overlay_page = overlay_reader.pages[0]
    page.merge_page(overlay_page)

    writer = PdfWriter()
    writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)

    return str(output_path)


def create_zip(data: TripReportData, output_zip: str | Path, report_docx: str, report_pdf: str, reason_docx: str | None = None) -> str:
    output_zip = Path(output_zip)
    output_zip.parent.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(output_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.write(report_docx, Path(report_docx).name)
        zf.write(report_pdf, Path(report_pdf).name)
        if reason_docx and Path(reason_docx).exists():
            zf.write(reason_docx, Path(reason_docx).name)

        zf.writestr("trip_report_data.json", json.dumps(data.to_dict(), ensure_ascii=False, indent=2))

        for p in collect_paths(data.to_dict()):
            path = Path(p)
            if path.exists():
                zf.write(path, f"attachments/{path.name}")

    return str(output_zip)
